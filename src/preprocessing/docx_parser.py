"""
DOCX Parser module for extracting text from Microsoft Word documents.
"""
import io
import os
import logging
from typing import Dict, Any, List, Optional, Union

import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

class DOCXParser:
    """
    Parser for extracting text and metadata from DOCX files.
    Uses python-docx for extraction.
    """

    def __init__(self, preserve_structure: bool = True):
        """
        Initialize the DOCX parser.
        
        Args:
            preserve_structure (bool): Whether to preserve document structure 
                                      (paragraphs, sections, tables, etc.)
        """
        self.preserve_structure = preserve_structure
        logger.info("Initialized DOCX parser with preserve_structure=%s", preserve_structure)

    def parse(self, file_path: Union[str, io.BytesIO]) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text and metadata.
        
        Args:
            file_path: Path to the DOCX file or BytesIO object
            
        Returns:
            Dict containing extracted text and metadata
        """
        logger.info("Parsing DOCX: %s", file_path if isinstance(file_path, str) else "BytesIO")
        
        try:
            # Load the document
            if isinstance(file_path, str):
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"DOCX file not found: {file_path}")
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract text content
            text_content = self._extract_text(doc)
            
            return {
                "metadata": metadata,
                "text": text_content,
                "structure": self._get_structure_info(doc) if self.preserve_structure else {},
                "success": True
            }
            
        except Exception as e:
            logger.error("Error parsing DOCX: %s", str(e), exc_info=True)
            return {
                "metadata": {},
                "text": "",
                "structure": {},
                "success": False,
                "error": str(e)
            }

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX document.
        
        Args:
            doc: Document object from python-docx
            
        Returns:
            Dict containing DOCX metadata
        """
        try:
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author,
                "created": core_props.created,
                "last_modified_by": core_props.last_modified_by,
                "modified": core_props.modified,
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            return {k: v for k, v in metadata.items() if v is not None}
        except Exception as e:
            logger.warning("Error extracting DOCX metadata: %s", str(e))
            return {"error": str(e)}

    def _extract_text(self, doc: Document) -> str:
        """
        Extract text from a DOCX document, preserving structure if required.
        
        Args:
            doc: Document object from python-docx
            
        Returns:
            Extracted text with structure preserved based on settings
        """
        if self.preserve_structure:
            # Extract with structure - iterate through document elements
            return self._extract_text_with_structure(doc)
        else:
            # Simple text extraction
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    def _extract_text_with_structure(self, doc: Document) -> str:
        """
        Extract text with structure preservation from a DOCX document.
        
        Args:
            doc: Document object from python-docx
            
        Returns:
            Text with preserved structure
        """
        text_content = []
        
        # Process all paragraphs and tables in the document in order
        for element in self._iter_block_items(doc):
            if isinstance(element, Paragraph):
                # Check if it's a heading
                if element.style.name.startswith('Heading'):
                    level = element.style.name.replace('Heading ', '')
                    text_content.append(f"{'#' * int(level) if level.isdigit() else '#'} {element.text}")
                else:
                    text_content.append(element.text)
            
            elif isinstance(element, Table):
                # Process table
                table_text = self._extract_table(element)
                text_content.append(table_text)
                
        return "\n\n".join([t for t in text_content if t.strip()])

    def _extract_table(self, table: Table) -> str:
        """
        Extract text from a table in a structured format.
        
        Args:
            table: Table object from python-docx
            
        Returns:
            Formatted table text
        """
        table_data = []
        
        # Extract header row
        header_row = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
        table_data.append(" | ".join(header_row))
        
        # Add separator
        table_data.append("-" * (len(" | ".join(header_row)) or 10))
        
        # Extract data rows
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(" | ".join(row_data))
            
        return "\n".join(table_data)

    def _get_structure_info(self, doc: Document) -> Dict[str, Any]:
        """
        Extract structural information from the document.
        
        Args:
            doc: Document object from python-docx
            
        Returns:
            Dict with structural information
        """
        # Count headings by level
        headings = {}
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                if level.isdigit():
                    headings[f"h{level}"] = headings.get(f"h{level}", 0) + 1
        
        return {
            "headings": headings,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }

    def _iter_block_items(self, doc: Document):
        """
        Iterate through all block items in the document in order,
        including paragraphs and tables.
        
        Args:
            doc: Document object from python-docx
            
        Yields:
            Paragraph and Table objects in document order
        """
        # Based on python-docx documentation for iterating through mixed content
        body = doc._body._body
        for child in body:
            if isinstance(child, CT_P):
                yield Paragraph(child, doc._body)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc._body)

    @staticmethod
    def is_valid_docx(file_path: Union[str, io.BytesIO]) -> bool:
        """
        Check if a file is a valid DOCX document.
        
        Args:
            file_path: Path to the file or BytesIO object
            
        Returns:
            Boolean indicating if the file is a valid DOCX
        """
        try:
            if isinstance(file_path, str):
                docx.Document(file_path)
            else:
                # Reset BytesIO position
                file_path.seek(0)
                docx.Document(file_path)
                # Reset again after checking
                file_path.seek(0)
            return True
        except Exception:
            return False